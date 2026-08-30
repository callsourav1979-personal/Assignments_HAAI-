"""
CV Sorting and Candidate Matching System

This application processes Job Descriptions and candidate CVs,
extracts structured information using a Large Language Model,
evaluates candidate skills, professional experience, and education,
and ranks candidates based on an overall matching score.
"""

#################################################################################################################
##   Imports & Device configuration
#################################################################################################################

"""
Imports the required Python libraries and configures the execution
environment for the CV Sorting and Candidate Matching application.

The imported libraries support the following capabilities:

    torch:
        Model execution and GPU/CPU device management.

    sys:
        Python runtime and version information.

    pypdf / PdfReader:
        Text extraction from PDF documents.

    docx / Document:
        Text extraction from Microsoft Word (DOCX) documents.

    pytesseract:
        Optical Character Recognition (OCR) for image-based documents.

    pdf2image:
        Conversion of PDF pages into images for OCR processing.

    json:
        Processing and validation of structured JSON data.

    pandas:
        Creation and management of candidate result tables and CSV output.

    pathlib:
        File and folder path management.

    re:
        Regular expression operations for text cleaning and matching.

    difflib.SequenceMatcher:
        Text similarity comparison where required.

The execution device is automatically selected based on CUDA availability:

    CUDA available:
        The application uses the GPU for faster LLM inference.

    CUDA unavailable:
        The application runs on the CPU.

The section also displays the runtime environment information, including
Python version, PyTorch version, CUDA availability, installed document
processing library versions, and GPU hardware details when available.
"""

import torch
import sys
import pypdf
import docx
import pytesseract
import json
import pandas as pd
from pdf2image import convert_from_path
from pathlib import Path
from pypdf import PdfReader
from docx import Document
import re


device = "cuda" if torch.cuda.is_available() else "cpu"

print("Using Device:" , device)

if torch.cuda.is_available():
    print("CUDA version:" , torch.version.cuda)
    print("GPU device name:" , torch.cuda.get_device_name(0))
    print("GPU Memory:" , round(torch.cuda.get_device_properties(0).total_memory/1024**3,2),"GB")
else:
  print("Running on CPU")

#################################################################################################################
## Model Loading ##
#################################################################################################################

"""
Loads the Qwen instruction-tuned language model and tokenizer for
LLM-based information extraction and semantic candidate evaluation.

The function selects the appropriate model precision and processing
device based on the available runtime environment.

When a CUDA-enabled GPU is available, the model is loaded using
float16 precision to improve memory efficiency and performance.

When GPU is not available, the model is loaded using float32 precision
and executed on the CPU.

Model Used:
    Qwen/Qwen2.5-1.5B-Instruct

Components:
    tokenizer: Converts text prompts into model-compatible tokens and
               decodes generated tokens into text responses.

    model1: Loaded causal language model used for Job Description
            extraction, CV extraction, and semantic experience matching.

Device Handling:
    CUDA: Uses GPU with torch.float16 precision.
    CPU: Uses CPU with torch.float32 precision.
"""

from transformers import AutoTokenizer,AutoModelForCausalLM

model_name1 = "Qwen/Qwen2.5-1.5B-Instruct"

tokenizer=AutoTokenizer.from_pretrained(model_name1,trust_remote_code=True)
if device == "cuda":
   model=AutoModelForCausalLM.from_pretrained(model_name1,
                                             torch_dtype=torch.float16,
                                             trust_remote_code=True
                                             )

   model1=model.to("cuda")
else:
   model=AutoModelForCausalLM.from_pretrained(model_name1,
                                               torch_dtype=torch.float32 ,
                                               trust_remote_code=True)
   model1=model.to("cpu")

print("Model loaded Successfully in :" , device)

#################################################################################################################
## Define Helper Functions
#################################################################################################################

def extract_text_from_file(file_path):
    """
    Extract text from PDF or DOCX files.

      For PDF:
          Extracts texts from all pages.

      For DOCX:
          Extracts texts from normal paragraphs and tables.

      Parameters :
           file_path(str): Path to the PDF or DOCX file.
      Returns :
           str: Extracted text from the file.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    extension = path.suffix.lower()
    #---------------------------------
    # PDF
    #---------------------------------
    if extension == '.pdf':
        reader = PdfReader(str(path))

        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
              pages.append(text)
        return '\n'.join(pages).strip()
    #---------------------------------------
    # DOCX
    #---------------------------------------
    elif extension == '.docx':
        document = Document(str(path))

        sections = []
        # Extract normal paragraphs
        for paragraph in document.paragraphs:
            text =  paragraph.text.strip()
            if text:
              sections.append(text)

        # Extract tables
        for table in document.tables:
            sections.append("\n[TABLE ]")
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                      row_cells.append(cell_text)
                #Combine cells in the same row
                if row_cells:
                   sections.append(" | ".join(row_cells))
            sections.append("[/TABLE]")
        #Combine paragraphs and tables
        extracted_text = "\n".join(sections).strip()

        return extracted_text

    else:
      raise ValueError(f"Unsupported file type: {extension}" "Only PDF and DOCX files are supported.")

def extract_text_from_pdf_ocr(file_path):
    """
    Extract text from scanned/image based PDF files using OCR

     Parameters :
           file_path(str): Path to the PDF file.
      Returns :
           str: OCR Extracted text from the file.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    extension = path.suffix.lower()
    if extension != '.pdf':
      raise ValueError(f"This function supports PDF files only.")

    # Convert PDF pages into images
    pages = convert_from_path(str(path), dpi=300)

    extracted_pages = []

    # Process each page
    for page_number , page_image in enumerate(pages, start =1):
        print(f"Processing page {page_number}/{len(pages)}")
        #Run OCR
        text = pytesseract.image_to_string(page_image,config="--psm 6")

        #Remove unnessary whitespace
        text = text.strip()

        #Store page seperately
        page_text = (f"\n---PAGE {page_number} ---\n" f"{(text)}")

    extracted_pages.append(page_text)

    # Combine all pages
    final_text = "\n".join(extracted_pages)

    return final_text.strip()

def extract_document_text(file_path):
  """
  Main document-extraction wrapper
  Automatically selects the appropriate extraction method based on the file type
  and available text.

    Parameters:
       file_path (str) : Path to the resume file

    Returns:
       str : Extracted text from the resume
  """
  path = Path(file_path)

  #1. Check whether the file exists
  if not path.exists():
    raise FileNotFoundError(f"File not found: {file_path}")

  #2. Check supported file types
  extension = path.suffix.lower()
  if extension not in ['.pdf' , '.docx']:
    raise ValueError(f"Unsupported file type: {extension}" "Only PDF and DOCX files are supported.")

  #3 Handle PDF
  if extension == '.pdf':
    print(f"\nProcessng PDF: {path.name}")

    #First try normal PDF text extraction
    text = extract_text_from_file(path)

    #Check whether meaningful text was extracted
    if text and len(text.strip()) >= 100 :
      print("Text layer detected. Using standard PDF extraction.")
      return text.strip()
    if len(text.strip()) < 100 :
      print("Little or no text detected . Swtching to OCR ...")
      text = extract_text_from_pdf_ocr(path)
      return text.strip()

  #4 Handle DOCX
  elif extension == '.docx':
    print(f"\nProcessing DOCX: {path.name}")
    text = extract_text_from_file(path)
    return text.strip()

def clean_and_validate_json(response):
  """
Cleans and validates JSON responses generated by the LLM.

The function removes optional Markdown code fences from the response,
strips unnecessary whitespace, and converts the cleaned JSON string
into a Python dictionary.

Parameters:
    response (str): Raw text response generated by the LLM.

Returns:
    dict: Validated JSON response converted into a Python dictionary.

Raises:
    json.JSONDecodeError: If the cleaned response is not valid JSON.
"""
  response = response.strip()
  # Remove opening markdown fence
  if response.startswith("```json"):
    response = response[len("```json"):].strip()
  elif response.startswith("```"):
    response = response[len("```"):].strip()

  # Remove closing markdown fence
  if response.endswith("```"):
    response = response[:-3].strip()

  # Validate JSON
  return json.loads(response)

def convert_jd_to_json(jd_text):
  """
Converts unstructured Job Description text into structured JSON format
using an LLM.

The function extracts job-related information including job title, required
skills, experience requirements, educational qualifications, and job
responsibilities explicitly mentioned in the Job Description.

Parameters:
    jd_text (str): Extracted raw text from a Job Description document.

Returns:
    dict: Structured JSON containing job title, skills, experience,
          education, and responsibilities.
"""

  jd_prompt = """
  DOCUMENT TYPE: JOB DESCRIPTION (JD)

  You are a precise recruitment information extraction assistant.

  Your task is to extract information ONLY from the provided JOB DESCRIPTION.
  Do not use outside knowledge. Do not infer, assume, invent, or fabricate information.

  IMPORTANT:
  This is a JOB DESCRIPTION, NOT a candidate CV.

  Return ONLY valid JSON.
  Do not return explanations, comments, markdown, keywords, or any text outside the JSON object.

  Use EXACTLY this JSON structure:

  {
    "job_title": "",
    "skills": [],
    "experience": [],
    "education": [],
    "responsibilities": []
  }

  FIELD DEFINITIONS:

  1. job_title
    Extract the exact job/position title stated in the JD.

  2. skills
    Extract technical skills, technologies, frameworks, platforms, tools,
    methodologies, architectural patterns, and certifications explicitly
    mentioned as required or preferred in the JD.

  Do not invent related technologies that are not explicitly mentioned.

  3. experience
    IMPORTANT: For a JOB DESCRIPTION, "experience" means
    EMPLOYER-STATED EXPERIENCE REQUIREMENTS.

  It does NOT mean candidate employment history.

  Extract experience requirements such as:
  - minimum total years of professional/software engineering experience
  - minimum years of architectural or technical leadership experience
  - required years of experience in a particular area
  - required experience with particular types of systems
  - required experience with methodologies or environments

  Preserve the meaning and wording of the JD as closely as possible.

  For example, if the JD says:
  "Minimum of 8+ years of total software engineering experience"

  then return:

  "experience": [
    "Minimum of 8+ years of total software engineering experience"
  ]

  If the JD says:
  "at least 3+ years acting in a dedicated Architectural or Tech Lead capacity"

  then return:

  "experience": [
    "At least 3+ years acting in a dedicated Architectural or Tech Lead capacity"
  ]

  NEVER convert an experience requirement into a fake employment-history object.

  DO NOT create fields such as:
  "title", "company", "location", "duration", or "description"
  inside the JD experience array.

  4. education
  Extract ONLY education requirements explicitly stated in the JD.

  For example, if the JD says:
  "Bachelor's or Master's degree in Computer Science, Software Engineering,
  or an equivalent technical field."

  return the relevant education requirement using the information actually
  present in the JD.

  DO NOT invent:
  - university names
  - graduation years
  - degree dates
  - fields of study not stated in the JD
  - candidate education details

  5. responsibilities
  Extract responsibilities, duties, activities, and expectations explicitly
  stated in the JD.

  Preserve the meaning of the JD.

  CRITICAL RULES:

  1. Extract information ONLY from the provided JD.
  2. Do NOT use outside knowledge.
  3. Do NOT infer or assume missing information.
  4. Do NOT invent companies, universities, candidates, dates, job histories,
     qualifications, or other information.
  5. Do NOT create candidate employment history from a JD.
  6. For a JD, the "experience" field contains EMPLOYER REQUIREMENTS,
     not candidate work history.
  7. For a JD, experience items must be strings, not employment-history objects.
  8. For a JD, do not create "Company A", "Company B", "University X",
     or similar placeholder/fabricated values.
  9. If information is not available, return [] for list fields and "" for
     the job_title field.
  10. Do not output values such as "Unknown", "Not specified",
     "Not mentioned", or "None".
  11. Do not include personal information that is not relevant to the
     requested extraction.
  12. Do not create a "keywords" field.
  13. Do not add any fields to the JSON structure.
  14. Return ONLY the JSON object..

  NOW EXTRACT THE INFORMATION FROM THIS JOB_DESCRIPTION:
  ----------------------------JOB DESCRIPTION START----------------------------------
  """ + jd_text + """

  ----------------------------JOB DESCRIPTION END----------------------------------

  """
  # Create system and user messages for the LLM
  messages = [
      {"role": "system", "content": "You are a precise recruitment assistant that extracts information from CV"},
      {"role": "user", "content": jd_prompt}
    ]

  # Convert messages into the model-specific chat format
  text= tokenizer.apply_chat_template(messages , tokenize=False , add_generation_prompt= True)

  # Convert text into model input tensors
  # and move tensors to the configured device
  inputs = tokenizer(text, return_tensors="pt").to(device)

  # Generate deterministic LLM response
  with torch.no_grad():
    outputs = model1.generate(**inputs,max_new_tokens=500,do_sample=False)

  # Decode only the newly generated tokens
  jd_response = tokenizer.decode(outputs[0][inputs["input_ids"].shape[1]:],skip_special_tokens=True)

  # Clean and validate the generated JSON response
  return clean_and_validate_json(jd_response)

def convert_cv_to_json(cv_text):
  """
Converts unstructured CV text into structured JSON format using an LLM.

The function extracts candidate information including professional title,
skills, work experience, education, responsibilities, projects, and work
activities explicitly mentioned in the CV.

Parameters:
    cv_text (str): Extracted raw text from a candidate CV document.

Returns:
    dict: Structured JSON containing candidate name, job title, skills,
          experience, education, and responsibilities.
"""

  cv_prompt = """
You are an expert recruitment assistant.

Your task is to analyze the following CV/Resume and convert information explicitly stated in it into a structured JSON format.

CV/RESUME:
""" + cv_text + """

Extract the following information:

1. Candidate Name
2. Job title or professional title, if explicitly stated
3. Skills explicitly mentioned anywhere in the CV
4. Work experience explicitly mentioned in the CV
5. Education/qualifications explicitly mentioned in the CV
6. Job responsibilities, duties, projects, or work activities explicitly mentioned in the CV


Return ONLY valid JSON in exactly this format:

{
  "candidate_name": "",
  "job_title": "",
  "skills": [],
  "experience": [],
  "education": [],
  "responsibilities": []
}


IMPORTANT RULES:

1. Extract information only from the CV. Do not use outside knowledge,
   assumptions, or infer skills that are not explicitly mentioned.

2. Extract skills comprehensively from ALL sections of the CV, including:
   - Technical Skills
   - Skills Summary
   - Work Experience
   - Projects
   - Responsibilities
   - Certifications
   - Tools and Technologies

3. "skills" must contain actual skills, abilities, knowledge, tools,
   technologies, software, programming languages, frameworks, platforms,
   methodologies, or competencies explicitly mentioned in the CV.

4. When a skill statement contains multiple individual technologies,
   extract both the broader skill and explicitly mentioned individual tools.

   Example:
   "Vector Databases (Pinecone, Milvus)"

   Extract:
   "Vector Databases"
   "Pinecone"
   "Milvus"

5. Do NOT infer skills based on the candidate's job title.

   For example, do not assume an AI Engineer knows Python unless Python
   is explicitly mentioned in the CV.

6. "experience" must contain explicitly stated work experience details.

7. "education" must contain explicitly stated educational qualifications,
   degrees, diplomas, certifications, or fields of study.

8. "responsibilities" must contain actual duties, projects, work activities,
   or responsibilities explicitly stated in the CV.

9. Do NOT convert responsibilities into skills unless a specific technology,
   tool, programming language, framework, or competency is explicitly
   mentioned.

10. Do NOT convert education into skills.

11. Do NOT infer experience, skills, or qualifications that are not explicitly
    stated.

12. If information is not present, return an empty list [].

13. NEVER output "None specified", "Not specified", "Not mentioned",
    "Unknown", or similar text. Use [] instead.

14. Do NOT add any fields that are not present in the JSON structure.

15. Do not provide explanations, comments, markdown, or text outside
    the JSON object.

16. Return ONLY valid JSON.

17. Stop generating immediately after the closing }.


OUTPUT LIMITS:

- Do not reproduce the CV verbatim.
- Summarize experience entries concisely.
- Maximum 30 skills.
- Maximum 5 education entries.
- Do not repeat skills.
- Do not repeat responsibilities.
- Do not copy entire paragraphs from the CV.
- Extract only information required by the JSON schema.
"""
    # Create system and user messages for the LLM
  messages = [
        {
            "role": "system",
            "content": "You are a precise recruitment assistant that extracts structured information from CVs."
        },
        {
            "role": "user",
            "content": cv_prompt
        }
    ]

  # Convert messages into the model-specific chat format
  text = tokenizer.apply_chat_template(messages,tokenize=False, add_generation_prompt=True)

  # Convert text into model input tensors
  # and move tensors to the configured device
  inputs = tokenizer(text,return_tensors="pt").to(device)

  # Generate deterministic LLM response
  with torch.no_grad():
        outputs = model1.generate(**inputs,max_new_tokens=1200,do_sample=False)

  # Decode only the newly generated tokens
  cv_response = tokenizer.decode(outputs[0][inputs["input_ids"].shape[1]:],skip_special_tokens=True)

  # Clean and validate the generated JSON response
  return clean_and_validate_json(cv_response)

def llm_match_candidate(jd, cv):
  """
    Uses an LLM to evaluate the semantic relevance of a candidate's
    professional experience against a Job Description.

    The evaluation considers job titles, work history, responsibilities,
    projects, technical context, domain relevance, and semantic
    equivalence.

    Parameters:
        jd (dict): Structured Job Description data.
        cv (dict): Structured candidate CV data.

    Returns:
        dict: JSON containing candidate name, experience match score,
              and evidence-based reasoning.
  """

  # Create prompt for LLM-based semantic experience evaluation

  matching_prompt = f"""
You are an expert recruitment and candidate experience evaluation assistant.

Your task is to evaluate how well a candidate's PROFESSIONAL EXPERIENCE
matches a Job Description.

Perform semantic evaluation. Do not rely only on exact keyword matching.

JOB DESCRIPTION JSON:
{json.dumps(jd, indent=2)}

CANDIDATE CV JSON:
{json.dumps(cv, indent=2)}

Evaluate ONLY the candidate's professional experience relevance.

Consider the following:

1. Job title relevance
   - Compare the candidate's previous job titles with the target job title.

2. Work experience relevance
   - Compare the candidate's work history with the experience requirements
     stated in the Job Description.

3. Responsibilities and project relevance
   - Compare the candidate's demonstrated responsibilities, projects,
     and work activities with the responsibilities in the Job Description.

4. Demonstrated technical and domain experience
   - Consider technologies, methodologies, and domains demonstrated
     through actual work experience and responsibilities.

5. Semantic equivalence
   - Treat different wording with substantially similar meaning as relevant
     experience.

   Examples:
   - RAG systems and hybrid semantic search may demonstrate experience
     relevant to semantic search frameworks.
   - Model quantization and latency reduction may demonstrate inference
     optimization experience.
   - Kubeflow, SageMaker, and ML CI/CD pipelines may demonstrate
     MLOps and production deployment experience.

SCORING GUIDELINES:

90-100:
Highly relevant professional experience with strong evidence of alignment
with nearly all major Job Description experience requirements and
responsibilities.

70-89:
Substantial relevant professional experience with strong alignment to
many Job Description requirements, but some gaps remain.

50-69:
Partially relevant experience with meaningful alignment, but significant
gaps exist.

30-49:
Limited relevant professional experience.

0-29:
Little or no relevant professional experience.

IMPORTANT RULES:

1. Use ONLY evidence present in the supplied Job Description JSON and
   Candidate CV JSON.

2. Evaluate actual demonstrated professional experience, including job
   history, responsibilities, projects, and work activities.

3. Do NOT rely solely on exact keyword matching.

4. Recognize semantically equivalent experience expressed using different
   terminology.

5. Do NOT assume experience with a technology merely because the candidate
   has a related job title or another related technology.

6. Do NOT evaluate education.

7. Do NOT calculate skill match scores.

8. Do NOT calculate an overall score.

9. Do NOT provide a hiring recommendation.

10. The score must reflect the overall semantic relevance of the candidate's
    demonstrated professional experience to the Job Description.

11. Keep the reasoning concise and evidence-based.

12. Do not treat generic business terminology as equivalent to technical
    machine learning terminology.

13. Terms such as "pipeline", "production", "deployment", "training",
    "model", or "architecture" must be interpreted according to their
    professional context.

14. Business sales pipelines, commercial production, employee training,
    product deployment, or business models must NOT be considered
    equivalent to machine learning pipelines, ML model deployment,
    ML model training, or ML architecture.

15. When evaluating semantic equivalence, consider the surrounding
    professional context, not just individual words.

16. A candidate whose experience is primarily in an unrelated domain,
    such as sales, finance, HR, or marketing, should receive a low
    experience score unless the CV contains clear evidence of relevant
    AI, machine learning, data science, or technical engineering work.

Return ONLY valid JSON in exactly this format:

{{
    "candidate_name": "",
    "experience_match_score": 0,
    "reasoning": ""
}}

FIELD DEFINITIONS:

- candidate_name:
  Candidate name from the CV.

- experience_match_score:
  Integer score from 0 to 100 representing the semantic relevance of
  the candidate's professional experience to the Job Description.

- reasoning:
  A concise evidence-based explanation for the experience score.
  Mention major areas of alignment and, where relevant, significant
  limitations.

Return ONLY the JSON object.
Do not return markdown, comments, or explanations outside the JSON.
Stop immediately after the closing }}.
"""
  # Create system and user messages for the LLM
  messages = [
        {
            "role": "system",
            "content": (
                "You are a precise recruitment assistant specializing in "
                "semantic evaluation of professional experience."
            )
        },
        {
            "role": "user",
            "content": matching_prompt
        }
    ]

  # Convert messages into the model-specific chat format
  text = tokenizer.apply_chat_template(messages,tokenize=False,add_generation_prompt=True
    )

  # Convert text into model input tensors
  # and move tensors to the configured device
  inputs = tokenizer(text,return_tensors="pt").to(device)

  # Generate deterministic LLM response
  with torch.no_grad():
        outputs = model1.generate( **inputs, max_new_tokens=500, do_sample=False
        )

  # Decode only the newly generated tokens
  response = tokenizer.decode(outputs[0][inputs["input_ids"].shape[1]:], skip_special_tokens=True
    )

  print("LLM Experience Matching Completed:")

   # Clean and validate the generated JSON response
  return clean_and_validate_json(response)

#################################################################################################################
## Define File Processing Functions
#################################################################################################################

def read_jd_from_folder(jd_folder):
    """
    Reads a Job Description file from a specified folder and converts
    the extracted document text into structured JSON format.

    The function searches for the first file matching the naming
    convention JD_*.pdf or JD_*.docx.

    Parameters:
        jd_folder (str): Path to the folder containing the
                         Job Description file.

    Returns:
        dict: Structured Job Description information.
    """

    # Convert folder path to a Path object
    jd_folder = Path(jd_folder)

    # Validate that the JD folder exists
    if not jd_folder.exists():
        raise FileNotFoundError(
            f"JD Folder not found: {jd_folder}"
        )

    # Search for supported JD files
    jd_files = (
        list(jd_folder.glob("JD_*.pdf")) +
        list(jd_folder.glob("JD_*.docx"))
    )

    # Raise an error if no JD file is found
    if not jd_files:
        raise FileNotFoundError(
            "No JD file found in the folder."
        )

    # Select the first matching JD file
    jd_file = jd_files[0]

    print(f"JD file found: {jd_file}")

    # Extract text from the JD document
    jd_text = extract_document_text(
        str(jd_file)
    )

    # Convert extracted JD text into structured JSON
    jd_json = convert_jd_to_json(
        jd_text
    )

    # Return structured Job Description data
    return jd_json

def extract_multiple_cvs(cv_folder):
    """
    Extract text from all supported CV files in a folder.

    Supported formats:
      - PDF
      - DOCX

    Parameters:
       cv_folder (str) : Path to the folder containing CVs

    Returns:
       dict: Dictionary containing filename and extracted text
    """

    folder = Path(cv_folder)
    if not folder.exists():
        raise FileNotFoundError(f"CV Folder not found: {cv_folder}")

    if not folder.is_dir():
        raise ValueError(f"Path is not a directory: {cv_folder}")

    # Find all PDF and DOCX files
    cv_files = sorted(
                       [ file
                         for file in folder.iterdir()
                         if file.is_file()
                         and file.name.startswith("CV_")
                         and file.suffix.lower() in [".pdf" , ".docx"]
                        ]
                      )
    if not cv_files:
      raise ValueError(f"No PDF or DOCX files found in :  {cv_folder}")

    cv_data = {}

    for cv_file in cv_files:
      print("="*60)
      print(f"Processing CV: {cv_file.name}")
      print("="*60)

      try:
        text = extract_document_text(cv_file)
        cv_data[cv_file.name] = text
        print(f"Characters Extracted: {len(text)}")
      except Exception as e:
        print(f"Error processing {cv_file.name}: {e}")
        cv_data[cv_file.name] = ""

    return cv_data

def process_all_cvs(cv_folder):
    """
    Processes all CV files in a specified folder and converts their
    extracted text into structured JSON format.

    Parameters:
        cv_folder (str): Path to the folder containing CV files.

    Returns:
        dict: Dictionary containing structured JSON data for each
              successfully processed CV, using filenames as keys.
    """

    # Extract text from all supported CV files
    all_cvs = extract_multiple_cvs(cv_folder)

    print(
        "Total CVs processed:",
        len(all_cvs)
    )

    # Initialize dictionary to store structured CV JSON
    all_cv_json = {}

    # Process each extracted CV
    for filename, cv_text in all_cvs.items():

        print(f"Processing: {filename}")

        try:
            # Convert extracted CV text into structured JSON
            cv_json = convert_cv_to_json(
                cv_text
            )

            # Store JSON using filename as the key
            all_cv_json[filename] = cv_json

            print("Successfully converted")

        except Exception as e:
            # Continue processing remaining CVs if one conversion fails
            print(
                f"Error processing {filename}: {e}"
            )

    print(
        "Total CVs converted:",
        len(all_cv_json)
    )

    return all_cv_json

#################################################################################################################
## Define Matching Functions
#################################################################################################################

# ---------------------------------------------------------
# 1. TEXT NORMALIZATION
# ---------------------------------------------------------

def normalize(text):
    """
    Standardizes text for consistent comparison.

    The function converts text to lowercase, removes unwanted special
    characters, normalizes whitespace, and removes leading/trailing spaces.

    Parameters:
        text: Input text to be normalized.

    Returns:
        str: Cleaned and standardized text.
    """

    # Return empty string for empty input
    if not text:
        return ""

    # Convert input to string and lowercase
    text = str(text).lower()

    # Remove unwanted characters while preserving common
    # technical symbols such as +, #, ., /, and -
    text = re.sub(
        r"[^a-z0-9+#./ -]",
        " ",
        text
    )

    # Replace multiple spaces and whitespace characters
    # with a single space
    text = re.sub(r"\s+", " ", text)

    # Remove leading and trailing spaces
    return text.strip()


# ---------------------------------------------------------
# 2. FLATTEN CV / JD SKILLS
# ---------------------------------------------------------

def get_skills(data):
    """
    Extracts and expands skills from structured JD or CV data.

    Skills containing comma-separated values inside parentheses are
    expanded into individual skills to improve matching accuracy.

    Parameters:
        data (dict): Structured JD or CV dictionary containing a
                     'skills' list.

    Returns:
        list: Expanded list of skills.
    """

    # Retrieve skills list from the input data
    skills = data.get("skills", [])

    # Initialize list to store expanded skills
    expanded_skills = []

    # Process each skill
    for skill in skills:

        # Add the original skill
        expanded_skills.append(skill.strip())

        # Extract text contained inside parentheses
        matches = re.findall(r'\((.*?)\)',skill)

        # Process each bracketed value
        for match in matches:

            # Split comma-separated sub-skills
            sub_skills = match.split(",")

            # Add each sub-skill individually
            for sub_skill in sub_skills:
                expanded_skills.append(
                    sub_skill.strip()
                )

    return expanded_skills


# ---------------------------------------------------------
# 3. DIRECT / FUZZY SKILL MATCHING
# ---------------------------------------------------------
def skill_matches(jd_skills, cv_skills):
    """
    Compares Job Description skills with candidate CV skills.

    Matching is case-insensitive and supports exact and partial
    string matching.

    Parameters:
        jd_skills (list): List of skills required in the Job Description.
        cv_skills (list): List of skills extracted from the candidate CV.

    Returns:
        tuple: Two lists containing matched skills and missing skills.
    """

    # Initialize lists for matched and missing JD skills
    matched_skills = []
    missing_skills = []

    # Normalize CV skills for case-insensitive comparison
    cv_skills_lower = [
        skill.lower().strip()
        for skill in cv_skills
    ]

    # Compare each Job Description skill
    for jd_skill in jd_skills:

        # Normalize JD skill
        jd_skill_lower = jd_skill.lower().strip()

        found = False

        # Compare against each CV skill
        for cv_skill in cv_skills_lower:

            # Check for exact or partial skill match
            if (
                jd_skill_lower == cv_skill
                or jd_skill_lower in cv_skill
                or cv_skill in jd_skill_lower
            ):
                found = True
                break

        # Categorize the JD skill based on matching result
        if found:
            matched_skills.append(jd_skill)
        else:
            missing_skills.append(jd_skill)

    return matched_skills, missing_skills

# ---------------------------------------------------------
# 4. EXPERIENCE MATCHING
# ---------------------------------------------------------

def calculate_experience_score(jd, cv):
    """
    Calculates semantic experience relevance using LLM2.

    Parameters:
        jd (dict): Structured Job Description JSON.
        cv (dict): Structured Candidate CV JSON.

    Returns:
        tuple: Experience match score and reasoning.
    """

    # Get semantic experience evaluation from LLM2
    llm_result = llm_match_candidate(jd, cv)

    # Handle empty or invalid LLM response
    if not llm_result:
        return 0, ""

    # Extract experience score
    experience_score = llm_result.get(
        "experience_match_score",
        0
    )

    # Extract LLM-generated reasoning
    experience_reasoning = llm_result.get(
        "reasoning",
        ""
    )

    return experience_score, experience_reasoning

# ---------------------------------------------------------
# 5. EDUCATION MATCHING
# ---------------------------------------------------------

def calculate_education_score(jd, cv):
    """
    Calculates the education match score between a Job Description
    and a candidate CV.

    The function evaluates educational field relevance and degree
    qualification level.

    Parameters:
        jd (dict): Structured Job Description data.
        cv (dict): Structured candidate CV data.

    Returns:
        int: Education match score between 0 and 100.
    """

    # Extract education requirements from JD and qualifications from CV
    jd_education = jd.get("education", [])
    cv_education = cv.get("education", [])

    # If JD does not specify education requirements,
    # do not penalize the candidate
    if not jd_education:
        return 100

    # Candidate has no education information
    if not cv_education:
        return 0

    # Convert JD education requirements into normalized text
    jd_text = normalize(" ".join(
        item if isinstance(item, str) else str(item)
        for item in jd_education
    ))

    # Extract relevant education details from the CV
    cv_parts = []

    for item in cv_education:

        if isinstance(item, dict):
            # Extract degree and field of study
            cv_parts.append(str(item.get("degree", "")))
            cv_parts.append(str(item.get("field_of_study", "")))
        else:
            cv_parts.append(str(item))

    # Convert CV education information into normalized text
    cv_text = normalize(" ".join(cv_parts))

    # --------------------------------
    # FIELD / SUBJECT MATCHING
    # --------------------------------

    # Define recognized technical education fields
    technical_fields = [
        "computer science",
        "software engineering",
        "computer engineering",
        "information technology",
        "information systems",
        "electrical engineering",
        "electronics",
        "engineering"
    ]

    # Identify technical fields required by the JD
    jd_fields = [
        field for field in technical_fields
        if field in jd_text
    ]

    # Identify technical fields present in the candidate CV
    cv_fields = [
        field for field in technical_fields
        if field in cv_text
    ]

    # If JD specifies a technical field,
    # evaluate the relevance of the candidate's field
    if jd_fields:

        # Candidate does not have a recognized technical field
        if not cv_fields:
            return 0

        # Candidate has a matching technical field
        if any(field in cv_fields for field in jd_fields):
            return 100

        # Candidate has a different technical field
        return 50

    # --------------------------------
    # DEGREE MATCHING
    # --------------------------------

    # Check Master's or MBA requirement
    if "master" in jd_text or "mba" in jd_text:
        if "master" in cv_text or "mba" in cv_text:
            return 100

    # Check Bachelor's degree requirement
    if "bachelor" in jd_text:
        if "bachelor" in cv_text or "master" in cv_text:
            return 100

    # Education requirement not matched
    return 0
# ---------------------------------------------------------
# 6. OVERALL MATCHING
# ---------------------------------------------------------

def match_candidate(jd, cv):
    """
    Performs complete candidate matching against a Job Description.

    The function calculates skill, experience, and education match
    scores, combines them using weighted scoring, and generates a
    final recommendation.

    Parameters:
        jd (dict): Structured Job Description data.
        cv (dict): Structured candidate CV data.

    Returns:
        dict: Complete candidate matching result including individual
              scores, overall score, matched skills, missing skills,
              experience reasoning, and recommendation.
    """

    # Extract and expand skills from the Job Description and CV
    jd_skills = get_skills(jd)
    cv_skills = get_skills(cv)

    # Identify matched and missing Job Description skills
    matched_skills, missing_skills = skill_matches(
        jd_skills,
        cv_skills
    )

    # Calculate skill match percentage
    if jd_skills:
        skill_score = round(
            len(matched_skills) /
            len(jd_skills) *
            100
        )
    else:
        skill_score = 0

    # Calculate semantic experience match score using LLM2
    experience_score, experience_reasoning = (
        calculate_experience_score(jd, cv)
    )

    # Calculate education qualification match score
    education_score = calculate_education_score(
        jd,
        cv
    )

    # Calculate weighted overall candidate score
    # Skills: 50%, Experience: 30%, Education: 20%
    overall_score = round(
        skill_score * 0.50 +
        experience_score * 0.30 +
        education_score * 0.20
    )

    # Generate recommendation based on overall score
    if overall_score >= 80:
        recommendation = "Strong Match"

    elif overall_score >= 60:
        recommendation = "Good Match"

    elif overall_score >= 40:
        recommendation = "Moderate Match"

    elif overall_score >= 20:
        recommendation = "Weak Match"

    else:
        recommendation = "Poor Match"

    # Extract candidate name
    candidate_name = cv.get(
        "candidate_name",
        ""
    )

    # Return complete candidate evaluation result
    return {
        "candidate_name": candidate_name,
        "overall_score": overall_score,
        "skill_match_score": skill_score,
        "experience_match_score": experience_score,
        "experience_reasoning": experience_reasoning,
        "education_match_score": education_score,
        "matched_skills": matched_skills,
        "missing_skills": missing_skills,
        "recommendation": recommendation
    }

def match_all_cvs(jd_json, all_cv_json):
    """
    Matches multiple candidate CVs against a Job Description.

    Parameters:
        jd_json (dict): Structured Job Description data.
        all_cv_json (dict): Dictionary containing structured CV data,
                            with filenames as keys.

    Returns:
        list: List of candidate matching result dictionaries.
    """

    # Initialize list to store matching results
    results = []

    # Process each CV individually
    for filename, cv_json in all_cv_json.items():

        print(f"Matching CV: {filename}")

        try:
            # Match candidate CV against the Job Description
            result = match_candidate(
                jd_json,
                cv_json
            )

            # Add source CV filename to the result
            result["cv_filename"] = filename

            # Store candidate result
            results.append(result)

        except Exception as e:
            # Continue processing remaining CVs if one CV fails
            print(
                f"Error matching {filename}: {e}"
            )

    return results

#################################################################################################################
## Define Ranking Function
#################################################################################################################
def rank_candidates(results):
    """
    Sorts candidates based on overall matching score and assigns ranks.

    Parameters:
        results (list): List of candidate matching result dictionaries.

    Returns:
        list: Candidate results sorted in descending order of overall score
              with a rank assigned to each candidate.
    """

    # Sort candidates by overall score in descending order
    ranked_results = sorted(
        results,
        key=lambda x: x.get("overall_score", 0),
        reverse=True
    )

    # Assign sequential rank starting from 1
    for rank, result in enumerate(ranked_results, start=1):
        result["rank"] = rank

    return ranked_results

#################################################################################################################
## Define Result Export Functions
#################################################################################################################

def create_result_table(ranked_results):
    """
    Creates a Pandas DataFrame from ranked candidate results.

    Parameters:
        ranked_results (list): List of ranked candidate dictionaries.

    Returns:
        pd.DataFrame: Formatted candidate ranking table.
    """

    # Convert ranked results into a DataFrame
    results_df = pd.DataFrame(ranked_results)

    # Select and arrange important columns if they exist
    preferred_columns = [
      "rank",
      "candidate_name",
      "cv_filename",
      "skill_match_score",
      "experience_match_score",
      "education_match_score",
      "overall_score",
      "recommendation"
    ]

    available_columns = [
        column for column in preferred_columns
        if column in results_df.columns
    ]

    # Keep preferred columns first
    remaining_columns = [
        column for column in results_df.columns
        if column not in available_columns
    ]

    results_df = results_df[
        available_columns + remaining_columns
    ]

    return results_df

def export_results(results_df, ranked_results, output_folder):
    """
    Exports candidate ranking results to CSV and JSON files.

    Parameters:
        results_df (pd.DataFrame): Candidate ranking table.
        ranked_results (list): Ranked candidate results.
        output_folder (str): Directory where output files will be saved.
    """

    # Convert output path to Path object
    output_path = Path(output_folder)

    # Create output directory if it does not exist
    output_path.mkdir(parents=True, exist_ok=True)

    # Define output file paths
    json_file = output_path / "candidate_ranking.json"
    csv_file = output_path / "candidate_ranking.csv"

    # Save ranked results as JSON
    with open(json_file, "w", encoding="utf-8") as f:
        json.dump(
            ranked_results,
            f,
            indent=2,
            ensure_ascii=False
        )

    # Save result table as CSV
    results_df.to_csv(
        csv_file,
        index=False,
        encoding="utf-8"
    )

    print("Results saved successfully.")
    print(f"JSON file: {json_file}")
    print(f"CSV file: {csv_file}")


#################################################################################################################
## Main Controller Function
#################################################################################################################


def run_cv_sorting(jd_folder, cv_folder, output_folder):
    """
    Executes the complete CV sorting workflow.

    Workflow:
        1. Read and convert Job Description to JSON.
        2. Read and convert all CVs to JSON.
        3. Match each CV against the Job Description.
        4. Rank candidates based on matching scores.
        5. Create a result table.
        6. Export results to CSV and JSON.

    Parameters:
        jd_folder (str): Path to the folder containing the Job Description.
        cv_folder (str): Path to the folder containing CV documents.
        output_folder (str): Path where output files will be saved.

    Returns:
        pd.DataFrame: Final ranked candidate result table.
    """

    print("=" * 60)
    print("STARTING CV SORTING PROCESS")
    print("=" * 60)

    # --------------------------------------------------
    # Step 1: Read and process Job Description
    # --------------------------------------------------
    print("\nStep 1: Processing Job Description...")

    jd_json = read_jd_from_folder(jd_folder)

    print("Job Description processed successfully.")

    # --------------------------------------------------
    # Step 2: Read and process all CVs
    # --------------------------------------------------
    print("\nStep 2: Processing CVs...")

    all_cv_json = process_all_cvs(cv_folder)

    print(f"Total CVs processed: {len(all_cv_json)}")

    # --------------------------------------------------
    # Step 3: Match CVs against Job Description
    # --------------------------------------------------
    print("\nStep 3: Matching candidates with Job Description...")

    results = match_all_cvs(jd_json, all_cv_json)

    print(f"Total candidates matched: {len(results)}")

    # --------------------------------------------------
    # Step 4: Rank candidates
    # --------------------------------------------------
    print("\nStep 4: Ranking candidates...")

    ranked_results = rank_candidates(results)

    print("Candidate ranking completed.")

    # --------------------------------------------------
    # Step 5: Create result table
    # --------------------------------------------------
    print("\nStep 5: Creating result table...")

    results_df = create_result_table(ranked_results)

    print("Result table created successfully.")

    # --------------------------------------------------
    # Step 6: Export results
    # --------------------------------------------------
    print("\nStep 6: Exporting results...")

    export_results(
        results_df,
        ranked_results,
        output_folder
    )

    print("\n" + "=" * 60)
    print("CV SORTING PROCESS COMPLETED SUCCESSFULLY")
    print("=" * 60)

    return ranked_results , results_df

#################################################################################################################
## Calling Function
#################################################################################################################
if __name__ == "__main__":

    import sys

    if len(sys.argv) != 4:
        print("Usage:")
        print("python cv_sorting.py <jd_folder> <cv_folder> <output_folder>")
        sys.exit(1)

    jd_folder = sys.argv[1]
    cv_folder = sys.argv[2]
    output_folder = sys.argv[3]

    run_cv_sorting(
        jd_folder=jd_folder,
        cv_folder=cv_folder,
        output_folder=output_folder
    )

